import re
import io
from pathlib import Path
from typing import Tuple, Optional
try:
    import pymupdf as fitz
except ImportError:
    import fitz
import docx
from backend.utils.config import MAX_FILE_SIZE_MB, ALLOWED_EXTENSIONS

class DocumentParsingError(Exception):
    pass

def normalize_text(text: str) -> str:
    """
    Normalizes raw extracted contract text:
    - Normalizes inconsistent whitespace & tabs
    - Fixes broken soft line breaks inside sentences
    - Removes repeated header/footer patterns
    - Collapses 3+ consecutive newlines to 2 newlines (standard paragraph breaks)
    - Trims leading/trailing whitespace
    """
    if not text:
        return ""

    # Replace carriage returns
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Replace multiple horizontal spaces/tabs with a single space
    text = re.sub(r"[ \t]+", " ", text)

    # Remove repeated page number lines e.g. "Page 1 of 5", "- 2 -", "Page 3"
    text = re.sub(r"(?i)\n\s*page\s+\d+(\s+of\s+\d+)?\s*\n", "\n\n", text)
    text = re.sub(r"\n\s*[-—]\s*\d+\s*[-—]\s*\n", "\n\n", text)

    # Join broken lines where a line ends in a lowercase letter/comma and the next begins with a lowercase letter
    text = re.sub(r"([a-z,;])\n([a-z])", r"\1 \2", text)

    # Collapse excessive newlines (keep at most 2 newlines for paragraph separation)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()

def extract_pdf_text(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts text from PDF bytes using PyMuPDF (fitz).
    Returns (extracted_text, page_count).
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = len(doc)
        if page_count == 0:
            raise DocumentParsingError("The PDF document contains 0 pages.")

        text_parts = []
        for page_num in range(page_count):
            page = doc.load_page(page_num)
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text.strip())

        full_text = "\n\n".join(text_parts)
        normalized = normalize_text(full_text)

        if not normalized:
            raise DocumentParsingError("No readable text found in the PDF. It may be a scanned image or protected.")

        return normalized, page_count
    except Exception as e:
        if isinstance(e, DocumentParsingError):
            raise e
        raise DocumentParsingError(f"Failed to parse PDF document: {str(e)}")

def extract_docx_text(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts text from DOCX bytes using python-docx.
    Returns (extracted_text, estimated_page_count).
    """
    try:
        file_stream = io.BytesIO(file_bytes)
        doc = docx.Document(file_stream)

        text_parts = []
        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n\n".join(text_parts)
        normalized = normalize_text(full_text)

        if not normalized:
            raise DocumentParsingError("No readable text found in the DOCX document.")

        # Estimate page count roughly (~500 words per page)
        words = len(normalized.split())
        estimated_pages = max(1, (words + 499) // 500)

        return normalized, estimated_pages
    except Exception as e:
        if isinstance(e, DocumentParsingError):
            raise e
        raise DocumentParsingError(f"Failed to parse DOCX document: {str(e)}")

def extract_txt_text(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extracts text from plain text bytes.
    """
    try:
        raw_text = file_bytes.decode("utf-8", errors="replace")
        normalized = normalize_text(raw_text)
        if not normalized:
            raise DocumentParsingError("The text document is empty.")
        words = len(normalized.split())
        estimated_pages = max(1, (words + 499) // 500)
        return normalized, estimated_pages
    except Exception as e:
        if isinstance(e, DocumentParsingError):
            raise e
        raise DocumentParsingError(f"Failed to parse text document: {str(e)}")

def validate_and_extract(filename: str, file_bytes: bytes) -> Tuple[str, str, int, int]:
    """
    Validates file format, size, and extracts text.
    Returns (normalized_text, file_type, page_count, char_count).
    """
    if not file_bytes or len(file_bytes) == 0:
        raise DocumentParsingError("The uploaded file is empty (0 bytes).")

    # Check file size
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise DocumentParsingError(f"File size exceeds maximum allowed limit of {MAX_FILE_SIZE_MB}MB.")

    # Determine extension
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise DocumentParsingError(f"Unsupported file format '{ext}'. Allowed formats: PDF (.pdf), Word (.docx), Text (.txt).")

    file_type = ext.replace(".", "")

    if ext == ".pdf":
        text, page_count = extract_pdf_text(file_bytes)
    elif ext == ".docx":
        text, page_count = extract_docx_text(file_bytes)
    elif ext == ".txt":
        text, page_count = extract_txt_text(file_bytes)
    else:
        raise DocumentParsingError(f"Unsupported file extension: {ext}")

    char_count = len(text)
    return text, file_type, page_count, char_count
